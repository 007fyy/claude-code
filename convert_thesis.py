#!/usr/bin/env python
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

MD_PATH = 'L:/claude-code/毕业论文_范盈盈_2022058207.md'
DOCX_PATH = 'L:/claude-code/毕业论文_范盈盈_2022058207.docx'

SZ_TITLE = 22; SZ_H1 = 16; SZ_H2 = 14; SZ_H3 = 12; SZ_BODY = 12; SZ_TABLE = 10.5

def set_cn_font(run, cn='宋体', en='Times New Roman', size=12, bold=False, italic=False):
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rFonts.set(qn(attr), en)
    rFonts.set(qn('w:eastAsia'), cn)

def set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5,
                 space_before=0, space_after=0, first_indent_pt=0, page_break_before=False):
    pf = para.paragraph_format
    pf.alignment = align
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_indent_pt:
        pf.first_line_indent = Pt(first_indent_pt)
    if page_break_before:
        pf.page_break_before = True

def add_inline_text(para, text, cn='宋体', en='Times New Roman', size=12, bold=False):
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        set_cn_font(run, cn=cn, en=en, size=size, bold=(bold or idx % 2 == 1))

def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def add_page_number(doc):
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    set_cn_font(run, size=10.5)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def is_table_row(line):
    s = line.strip()
    return s.startswith('|') and s.endswith('|')

def is_table_sep(line):
    return bool(re.match(r'^\s*\|[-| :]+\|\s*$', line))

def parse_table_row(line):
    cells = line.strip().strip('|').split('|')
    return [c.strip() for c in cells]

def flush_table(doc, table_rows):
    if not table_rows:
        return
    ncols = len(table_rows[0])
    table = doc.add_table(rows=len(table_rows), cols=ncols)
    table.style = 'Table Grid'
    for r, row_data in enumerate(table_rows):
        for c, cell_text in enumerate(row_data):
            if c >= len(table.rows[r].cells):
                continue
            cell = table.rows[r].cells[c]
            cell.text = ''
            para = cell.paragraphs[0]
            set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
                         space_before=0, space_after=0)
            is_header = (r == 0)
            add_inline_text(para, cell_text, cn='宋体', size=SZ_TABLE, bold=is_header)
    add_table_borders(table)
    # Space after table
    after = doc.add_paragraph()
    set_para_fmt(after, space_before=0, space_after=6)

def convert():
    with open(MD_PATH, encoding='utf-8') as f:
        lines = [l.rstrip('\n') for l in f.readlines()]

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    add_page_number(doc)

    i = 0
    cover_done = False
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Flush pending table when we leave table context ──────────────────
        if in_table and not is_table_row(line):
            flush_table(doc, table_rows)
            table_rows = []
            in_table = False

        # ── Cover page (before first ---) ────────────────────────────────────
        if not cover_done:
            if stripped == '---':
                cover_done = True
                para = doc.add_paragraph()
                run = para.add_run()
                run.add_break(WD_BREAK.PAGE)
                i += 1
                continue
            if stripped.startswith('# '):
                para = doc.add_paragraph()
                set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER,
                             line_spacing=1.5, space_before=72, space_after=24)
                run = para.add_run(stripped[2:])
                set_cn_font(run, cn='黑体', en='SimHei', size=SZ_TITLE, bold=True)
            elif stripped.startswith('**') and stripped.endswith('  '):
                para = doc.add_paragraph()
                set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
                add_inline_text(para, stripped.rstrip(), cn='宋体', size=SZ_BODY)
            elif stripped.startswith('**'):
                para = doc.add_paragraph()
                set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
                add_inline_text(para, stripped, cn='宋体', size=SZ_BODY)
            i += 1
            continue

        # ── Horizontal rule → page break ─────────────────────────────────────
        if stripped == '---':
            para = doc.add_paragraph()
            run = para.add_run()
            run.add_break(WD_BREAK.PAGE)
            i += 1
            continue

        # ── Empty line ────────────────────────────────────────────────────────
        if stripped == '':
            i += 1
            continue

        # ── Table row ─────────────────────────────────────────────────────────
        if is_table_row(line):
            if not is_table_sep(line):
                table_rows.append(parse_table_row(line))
            in_table = True
            i += 1
            continue

        # ── H2: chapter / abstract / references ──────────────────────────────
        if stripped.startswith('## '):
            text = stripped[3:]
            para = doc.add_paragraph()
            set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER,
                         line_spacing=1.5, space_before=12, space_after=12)
            run = para.add_run(text)
            set_cn_font(run, cn='黑体', en='SimHei', size=SZ_H1, bold=True)
            i += 1
            continue

        # ── H3: section ───────────────────────────────────────────────────────
        if stripped.startswith('### '):
            text = stripped[4:]
            para = doc.add_paragraph()
            set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT,
                         line_spacing=1.5, space_before=6, space_after=6)
            run = para.add_run(text)
            set_cn_font(run, cn='黑体', en='SimHei', size=SZ_H2, bold=True)
            i += 1
            continue

        # ── H4: subsection ────────────────────────────────────────────────────
        if stripped.startswith('#### '):
            text = stripped[5:]
            para = doc.add_paragraph()
            set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT,
                         line_spacing=1.5, space_before=3, space_after=3)
            run = para.add_run(text)
            set_cn_font(run, cn='黑体', en='SimHei', size=SZ_H3, bold=True)
            i += 1
            continue

        # ── Table caption: **表 X-X　...** ───────────────────────────────────
        if re.match(r'^\*\*表\s*\d', stripped):
            para = doc.add_paragraph()
            set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER,
                         line_spacing=1.5, space_before=6, space_after=3)
            caption = stripped.strip('*').strip()
            run = para.add_run(caption)
            set_cn_font(run, cn='宋体', size=SZ_BODY, bold=True)
            i += 1
            continue

        # ── Reference item: [N] ... ───────────────────────────────────────────
        if re.match(r'^\[\d+\]', stripped):
            para = doc.add_paragraph()
            set_para_fmt(para, line_spacing=1.5, space_before=0, space_after=0,
                         first_indent_pt=0)
            para.paragraph_format.left_indent = Pt(0)
            add_inline_text(para, stripped, cn='宋体', size=SZ_BODY)
            i += 1
            continue

        # ── Body text ─────────────────────────────────────────────────────────
        para = doc.add_paragraph()
        set_para_fmt(para, line_spacing=1.5, space_before=0, space_after=0,
                     first_indent_pt=SZ_BODY * 2)
        add_inline_text(para, stripped, cn='宋体', size=SZ_BODY)
        i += 1

    # Flush any remaining table
    if table_rows:
        flush_table(doc, table_rows)

    doc.save(DOCX_PATH)
    print(f'Saved: {DOCX_PATH}')

convert()
